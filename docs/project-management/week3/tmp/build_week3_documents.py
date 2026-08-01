from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, landscape
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfgen import canvas


ROOT = Path(r"C:\Users\Houro\Desktop\IT Elect 4\storytale")
OUT = ROOT / "docs" / "project-management" / "week3"
OUT.mkdir(parents=True, exist_ok=True)

PURPLE = "4C2BBF"
DARK = "171428"
GRAY = "5F5A6B"
LIGHT = "F3F0FA"
LAVENDER = "E8E0FF"
WHITE = "FFFFFF"
RED = "C2415A"
AMBER = "D88B22"
GREEN = "2F7A5A"


def pdf_color(value: str):
    return colors.HexColor(f"#{value}")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=85, start=110, bottom=85, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_width(cell, width_inches: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def add_bottom_border(paragraph, color=PURPLE, size=18) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "7")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_page_number(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("StoryTale  |  ")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(GRAY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def set_doc_defaults(doc: Document, title: str) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.58)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)
    set_page_number(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.05

    for style_name, size, color in (
        ("Title", 24, PURPLE),
        ("Heading 1", 16, PURPLE),
        ("Heading 2", 12, DARK),
    ):
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(5)

    props = doc.core_properties
    props.title = title
    props.subject = "IT Elect 4 - StoryTale Project Management"
    props.author = "John Benedict S. Alejo"
    props.keywords = "StoryTale, project management, timeline, milestones, risk plan"


def add_masthead(doc: Document, document_title: str, subtitle: str) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    left, right = table.rows[0].cells
    set_cell_width(left, 5.25)
    set_cell_width(right, 1.55)
    set_cell_shading(left, PURPLE)
    set_cell_shading(right, DARK)
    for cell in (left, right):
        set_cell_margins(cell, top=150, start=160, bottom=140, end=160)

    p = left.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run("STORYTALE")
    r.font.name = "Arial"
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(LAVENDER)

    p = left.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(document_title)
    r.font.name = "Arial"
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(WHITE)

    p = left.add_paragraph(subtitle)
    p.paragraph_format.space_after = Pt(0)
    for r in p.runs:
        r.font.name = "Arial"
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string(LAVENDER)

    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("IT ELECT 4")
    r.font.name = "Arial"
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(WHITE)
    p = right.add_paragraph("3 AUG 2026")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in p.runs:
        r.font.name = "Arial"
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string(LAVENDER)
    p = right.add_paragraph("John Benedict S. Alejo")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in p.runs:
        r.font.name = "Arial"
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string(WHITE)

    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def format_table(table, widths, header=True, body_font=8.6) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row_idx, row in enumerate(table.rows):
        keep_row_together(row)
        if row_idx == 0 and header:
            set_repeat_table_header(row)
        for idx, cell in enumerate(row.cells):
            if idx < len(widths):
                set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_idx == 0 and header:
                set_cell_shading(cell, PURPLE)
            elif row_idx % 2 == 0:
                set_cell_shading(cell, LIGHT)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(body_font)
                    if row_idx == 0 and header:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor.from_string(WHITE)
                    else:
                        run.font.color.rgb = RGBColor.from_string(DARK)


def add_summary_band(doc: Document, items) -> None:
    table = doc.add_table(rows=1, cols=len(items))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    width = 6.95 / len(items)
    for cell, (number, label) in zip(table.rows[0].cells, items):
        set_cell_width(cell, width)
        set_cell_shading(cell, LIGHT)
        set_cell_margins(cell, top=120, start=120, bottom=110, end=120)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(number)
        r.font.name = "Arial"
        r.font.size = Pt(16)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(PURPLE)
        p = cell.add_paragraph(label)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        for r in p.runs:
            r.font.name = "Arial"
            r.font.size = Pt(8)
            r.font.color.rgb = RGBColor.from_string(GRAY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


MILESTONES = [
    ("Proposal and wireframes approved", "Approved proposal and 3-5 screen user flow", "Week 1", "26 Jul 2026"),
    ("Functional Flutter foundation", "Organized app shell, theme, reusable navigation, and placeholder screens", "Week 2", "2 Aug 2026"),
    ("Local EPUB reader complete", "EPUB import, chapter list, reading view, and saved progress", "Week 4", "16 Aug 2026"),
    ("Translation and audio prototype", "DeepL Filipino translation and offline voice playback prototype", "Week 6", "30 Aug 2026"),
    ("Story analysis and asset catalogs", "Validated story data, reusable character, location, and foreground catalogs", "Week 8", "13 Sep 2026"),
    ("Animated Story Mode prototype", "Chapter-based visual novel playback with sprites, subtitles, camera, and movement", "Week 9", "20 Sep 2026"),
    ("System testing complete", "Test report covering mobile behavior, performance, data, APIs, and accessibility", "Week 10", "27 Sep 2026"),
    ("Final system and presentation", "Completed application, documentation, repository, and final presentation", "Week 11", "4 Oct 2026"),
]


def build_milestones_docx() -> Path:
    doc = Document()
    set_doc_defaults(doc, "StoryTale Milestones and Deliverables")
    add_masthead(
        doc,
        "Milestones and Deliverables",
        "A concise delivery map for the 11-week StoryTale implementation.",
    )

    p = doc.add_paragraph("Project Delivery Summary", style="Heading 1")
    add_bottom_border(p)
    doc.add_paragraph(
        "StoryTale is a local-first Flutter e-book library for students and English learners. "
        "It combines EPUB reading, English-to-Filipino translation, offline narration, and a "
        "chapter-based Animated Story Mode. The schedule prioritizes a usable reading experience "
        "before the AI-assisted visual features."
    )
    add_summary_band(
        doc,
        [
            ("11", "project weeks"),
            ("8", "major milestones"),
            ("4", "core reading modes"),
            ("1", "final integrated app"),
        ],
    )

    p = doc.add_paragraph("Major Milestones", style="Heading 1")
    add_bottom_border(p)
    table = doc.add_table(rows=1, cols=4)
    headers = ["Milestone", "Expected Output", "Target", "Target Date"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
    for milestone, output, target, target_date in MILESTONES:
        cells = table.add_row().cells
        for cell, text in zip(cells, (milestone, output, target, target_date)):
            cell.text = text
    format_table(table, [2.02, 3.16, 0.72, 1.05], body_font=8.2)

    doc.add_page_break()
    p = doc.add_paragraph("Deliverable Acceptance Criteria", style="Heading 1")
    add_bottom_border(p)
    criteria = [
        ("Functional", "The required flow can be completed on the target device without a blocking error."),
        ("Consistent", "Imported books, characters, locations, and generated assets keep stable identities across chapters."),
        ("Local-first", "Books, progress, translations, voice choices, and generated story packages remain available on the device."),
        ("Documented", "Architecture decisions, setup steps, limitations, tests, and final handoff are recorded."),
        ("Presentable", "The interface follows the approved StoryTale visual direction and is readable on mobile screens."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Criterion"
    table.rows[0].cells[1].text = "Definition of Done"
    for name, definition in criteria:
        cells = table.add_row().cells
        cells[0].text = name
        cells[1].text = definition
    format_table(table, [1.25, 5.7], body_font=8.8)

    p = doc.add_paragraph("Progress Reporting", style="Heading 1")
    add_bottom_border(p)
    doc.add_paragraph(
        "Progress will be checked at the end of each week using the working Flutter build, "
        "the related document or test evidence, and the next week's entry criteria. A milestone "
        "is marked complete only when its output is usable and recorded, not merely started."
    )

    path = OUT / "StoryTale_Milestones_and_Deliverables.docx"
    doc.save(path)
    return path


RISKS = [
    ("R1", "Cloud API quota, outage, or slow response", "Medium", "High", "Cache successful results; avoid duplicate calls; add retry and timeout rules.", "Keep the last valid result and allow the user to retry later.", "API error or quota warning"),
    ("R2", "Gemini story analysis returns inconsistent data", "Medium", "High", "Use strict structured output, validation, stable IDs, and chapter-to-book context.", "Reject invalid output and use a safe minimal chapter package.", "Schema or continuity check fails"),
    ("R3", "Generated characters or assets change appearance", "High", "High", "Lock reference images, character descriptions, reusable asset IDs, and approved catalogs.", "Regenerate only the failed variant or replace it with an approved placeholder.", "Identity score or visual review fails"),
    ("R4", "Offline voice packs make the app too large", "High", "Medium", "Limit the first release to five optional, quantized ONNX voice packs.", "Ship one default voice and let users download or remove other packs.", "Install size exceeds target"),
    ("R5", "Animated Story Mode performs poorly on mobile", "Medium", "High", "Pre-generate assets, cache audio, limit visible sprites, and profile on Android.", "Reduce motion, image resolution, and active layers while keeping subtitles and audio.", "Dropped frames or high memory use"),
    ("R6", "Different EPUB structures break import or chapters", "High", "High", "Normalize spine and HTML content; test multiple EPUB layouts and encodings.", "Show manual chapter review and keep the original EPUB available.", "Missing or merged chapters"),
    ("R7", "DeepL character allowance is consumed too quickly", "Medium", "Medium", "Translate per chapter, cache every result, and display usage before a request.", "Allow English-only reading until the allowance resets or is upgraded.", "Usage reaches 80 percent"),
    ("R8", "Local data or generated assets are lost", "Medium", "High", "Use durable local database and file storage with atomic saves and migrations.", "Restore from a local export or rebuild only the affected derived assets.", "Missing record after restart"),
    ("R9", "Schedule slips because the scope expands", "Medium", "High", "Protect the MVP, review weekly priorities, and defer music and advanced polish.", "Freeze new features and complete the reader, translation, audio, and one chapter demo.", "A milestone misses its week"),
    ("R10", "Copyright, privacy, or API key exposure", "Low", "High", "Use user-owned or public-domain EPUBs, local files, and server-side secrets.", "Remove exposed credentials, rotate keys, and disable affected generated content.", "Secret scan or content complaint"),
]


def build_risk_docx() -> Path:
    doc = Document()
    set_doc_defaults(doc, "StoryTale Risk Assessment and Contingency Plan")
    add_masthead(
        doc,
        "Risk Assessment & Contingency Plan",
        "Prevention, triggers, and recovery actions for the StoryTale implementation.",
    )

    p = doc.add_paragraph("Assessment Approach", style="Heading 1")
    add_bottom_border(p)
    doc.add_paragraph(
        "Risks are evaluated using qualitative probability and impact levels. High-impact risks "
        "receive an explicit prevention action, a measurable warning trigger, and a contingency "
        "that protects the core local reading experience."
    )
    add_summary_band(
        doc,
        [
            ("10", "tracked risks"),
            ("7", "high-impact risks"),
            ("3", "AI/API risks"),
            ("Weekly", "review frequency"),
        ],
    )

    p = doc.add_paragraph("Risk Priority Matrix", style="Heading 1")
    add_bottom_border(p)
    matrix = [
        ("High probability", "R4, R7", "R3, R6"),
        ("Medium probability", "-", "R1, R2, R5, R8, R9"),
        ("Low probability", "-", "R10"),
    ]
    table = doc.add_table(rows=1, cols=3)
    for cell, text in zip(table.rows[0].cells, ("Probability", "Medium Impact", "High Impact")):
        cell.text = text
    for probability, medium_items, high_items in matrix:
        cells = table.add_row().cells
        cells[0].text = probability
        cells[1].text = medium_items
        cells[2].text = high_items
    format_table(table, [1.65, 2.6, 2.7], body_font=9)
    set_cell_shading(table.rows[1].cells[1], AMBER)
    set_cell_shading(table.rows[1].cells[2], RED)
    set_cell_shading(table.rows[2].cells[2], AMBER)
    set_cell_shading(table.rows[3].cells[2], LIGHT)
    for row in table.rows[1:]:
        for cell in row.cells[1:]:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if cell._tc.get_or_add_tcPr().find(qn("w:shd")) is not None and cell._tc.get_or_add_tcPr().find(qn("w:shd")).get(qn("w:fill")) == RED:
                        run.font.color.rgb = RGBColor.from_string(WHITE)

    p = doc.add_paragraph("Risk Register and Contingencies", style="Heading 1")
    add_bottom_border(p)
    table = doc.add_table(rows=1, cols=7)
    headers = ["ID", "Risk", "Prob.", "Impact", "Prevention / Mitigation", "Contingency", "Trigger"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
    for item in RISKS:
        cells = table.add_row().cells
        for cell, text in zip(cells, item):
            cell.text = text
    format_table(table, [0.38, 1.18, 0.52, 0.52, 1.75, 1.55, 1.05], body_font=7.2)

    p = doc.add_paragraph("Response Rules", style="Heading 1")
    add_bottom_border(p)
    rules = [
        "Stop or isolate any change that threatens book data, API keys, or the ability to read an imported EPUB.",
        "Prefer a smaller working mode over a visually richer mode that crashes or blocks the reader.",
        "Record the trigger, decision, owner, and outcome in the weekly progress notes.",
        "Escalate any high-impact risk that remains open for two consecutive weekly reviews.",
        "Reassess probability and impact after mitigation is implemented.",
    ]
    for rule in rules:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(rule)

    p = doc.add_paragraph("Ownership and Review", style="Heading 1")
    add_bottom_border(p)
    doc.add_paragraph(
        "The student developer owns the register and reviews it every week. Technical risks are "
        "checked during implementation and targeted validation; schedule and scope risks are checked "
        "against the Gantt chart. Any activated contingency becomes the next priority until the "
        "core StoryTale reading flow is stable again."
    )

    path = OUT / "StoryTale_Risk_Assessment_and_Contingency_Plan.docx"
    doc.save(path)
    return path


def draw_wrapped(c: canvas.Canvas, text: str, x: float, y: float, max_width: float, font: str, size: float, leading: float, color=colors.black) -> float:
    c.setFillColor(color)
    c.setFont(font, size)
    words = text.split()
    line = ""
    lines = []
    for word in words:
        candidate = word if not line else f"{line} {word}"
        if stringWidth(candidate, font, size) <= max_width:
            line = candidate
        else:
            lines.append(line)
            line = word
    if line:
        lines.append(line)
    for item in lines:
        c.drawString(x, y, item)
        y -= leading
    return y


GANTT_ROWS = [
    ("PHASE 1 - Initiation", 1, 1, True),
    ("Proposal, audience, scope, wireframes", 1, 1, False),
    ("PHASE 2 - Foundation", 2, 2, True),
    ("Flutter structure, theme, navigation", 2, 2, False),
    ("Architecture and API decisions", 2, 2, False),
    ("PHASE 3 - Core Library", 3, 4, True),
    ("Project plan, local data design", 3, 3, False),
    ("EPUB import, chapters, reader progress", 3, 4, False),
    ("PHASE 4 - Language & Audio", 5, 6, True),
    ("DeepL translation and local cache", 5, 5, False),
    ("Offline TTS, voices, audiobook sync", 6, 6, False),
    ("PHASE 5 - Story Intelligence", 7, 8, True),
    ("Gemini analysis and Story Bible", 7, 7, False),
    ("Sprites, backgrounds, asset catalogs", 8, 8, False),
    ("PHASE 6 - Animated Story Mode", 9, 9, True),
    ("Shots, poses, camera, motion, subtitles", 9, 9, False),
    ("PHASE 7 - QA & Closure", 10, 11, True),
    ("Testing, optimization, risk closure", 10, 10, False),
    ("Final system, docs, presentation", 11, 11, False),
]


def build_gantt_pdf() -> Path:
    path = OUT / "StoryTale_Project_Timeline_Gantt.pdf"
    page_w, page_h = landscape(letter)
    c = canvas.Canvas(str(path), pagesize=(page_w, page_h))
    c.setTitle("StoryTale Project Timeline and Gantt Chart")
    c.setAuthor("John Benedict S. Alejo")

    margin = 28
    c.setFillColor(pdf_color(DARK))
    c.rect(0, page_h - 74, page_w, 74, fill=1, stroke=0)
    c.setFillColor(pdf_color(LAVENDER))
    c.setFont("Helvetica-Bold", 9)
    c.drawString(margin, page_h - 24, "STORYTALE  |  IT ELECT 4")
    c.setFillColor(colors.white)
    c.setFont("Helvetica-Bold", 23)
    c.drawString(margin, page_h - 52, "Project Timeline / Gantt Chart")
    c.setFont("Helvetica", 8.5)
    c.setFillColor(pdf_color(LAVENDER))
    c.drawRightString(page_w - margin, page_h - 28, "20 Jul - 4 Oct 2026")
    c.drawRightString(page_w - margin, page_h - 45, "Prepared by John Benedict S. Alejo")

    table_top = page_h - 97
    row_h = 21.8
    task_w = 258
    grid_x = margin + task_w
    grid_w = page_w - margin * 2 - task_w
    week_w = grid_w / 11

    c.setFillColor(pdf_color(PURPLE))
    c.rect(margin, table_top - 31, task_w, 31, fill=1, stroke=0)
    c.setFillColor(pdf_color(PURPLE))
    c.rect(grid_x, table_top - 31, grid_w, 31, fill=1, stroke=0)
    c.setFillColor(colors.white)
    c.setFont("Helvetica-Bold", 8.5)
    c.drawString(margin + 8, table_top - 19, "PROJECT PHASE / TASK")

    week_dates = [
        "20 Jul", "27 Jul", "3 Aug", "10 Aug", "17 Aug", "24 Aug",
        "31 Aug", "7 Sep", "14 Sep", "21 Sep", "28 Sep",
    ]
    for idx in range(11):
        x = grid_x + idx * week_w
        c.setStrokeColor(pdf_color("765DD0"))
        c.line(x, table_top, x, table_top - 31)
        c.setFillColor(colors.white)
        c.setFont("Helvetica-Bold", 7.2)
        c.drawCentredString(x + week_w / 2, table_top - 12, f"W{idx + 1}")
        c.setFont("Helvetica", 6.2)
        c.drawCentredString(x + week_w / 2, table_top - 23, week_dates[idx])
    c.line(grid_x + grid_w, table_top, grid_x + grid_w, table_top - 31)

    y = table_top - 31
    phase_color_cycle = [
        pdf_color("4C2BBF"),
        pdf_color("6747D1"),
        pdf_color("7A5ADD"),
        pdf_color("5E49B7"),
        pdf_color("3F6FB7"),
        pdf_color("5B48C8"),
        pdf_color("3B6C62"),
    ]
    phase_index = -1
    for row_idx, (label, start_week, end_week, is_phase) in enumerate(GANTT_ROWS):
        y -= row_h
        if is_phase:
            phase_index += 1
            c.setFillColor(pdf_color("EEEAF8"))
        else:
            c.setFillColor(colors.white if row_idx % 2 else pdf_color("FAF9FC"))
        c.rect(margin, y, task_w + grid_w, row_h, fill=1, stroke=0)

        c.setStrokeColor(pdf_color("D9D4E4"))
        c.setLineWidth(0.45)
        c.line(margin, y, page_w - margin, y)
        c.line(grid_x, y, grid_x, y + row_h)
        for idx in range(12):
            x = grid_x + idx * week_w
            c.line(x, y, x, y + row_h)

        c.setFillColor(pdf_color(DARK if is_phase else "403A4B"))
        c.setFont("Helvetica-Bold" if is_phase else "Helvetica", 7.4 if is_phase else 7.2)
        prefix = "" if is_phase else "  "
        c.drawString(margin + 7, y + 7.2, prefix + label)

        bar_x = grid_x + (start_week - 1) * week_w + 3
        bar_w = (end_week - start_week + 1) * week_w - 6
        bar_y = y + (5.2 if is_phase else 6.5)
        bar_h = 11.5 if is_phase else 8.7
        c.setFillColor(phase_color_cycle[phase_index])
        c.roundRect(bar_x, bar_y, bar_w, bar_h, 3, fill=1, stroke=0)

    c.setStrokeColor(pdf_color("B9B1CA"))
    c.line(margin, table_top, page_w - margin, table_top)
    c.line(margin, y, page_w - margin, y)

    legend_y = 35
    c.setFillColor(pdf_color(GRAY))
    c.setFont("Helvetica", 7.2)
    c.drawString(margin, legend_y + 9, "Duration: 11 weeks  |  Major milestones: 8  |  Final completion: 4 Oct 2026")
    c.setFillColor(pdf_color(PURPLE))
    c.roundRect(page_w - 218, legend_y + 4, 12, 8, 2, fill=1, stroke=0)
    c.setFillColor(pdf_color(GRAY))
    c.drawString(page_w - 201, legend_y + 5, "planned activity")
    c.setFillColor(pdf_color(LIGHT))
    c.rect(page_w - 125, legend_y + 4, 12, 8, fill=1, stroke=0)
    c.setFillColor(pdf_color(GRAY))
    c.drawString(page_w - 108, legend_y + 5, "phase row")

    c.save()
    return path


if __name__ == "__main__":
    outputs = [
        build_gantt_pdf(),
        build_milestones_docx(),
        build_risk_docx(),
    ]
    for output in outputs:
        print(output)
